# Framework supporting MLOps Apps
import streamlit as st 
# Large Language Model Library
from langchain.llms import OpenAI
import docx
import datetime
import time
import base64

# Front-End
with st.sidebar:
    st.markdown("<h1 style='text-align:center;font-family:Georgia'>🩺 AIVABOT </h1>", unsafe_allow_html=True)
    st.markdown("Welcome to AIVABOT, your Medical Summary Report Generator and Chat Bot. AIVABOT is designed to help you generate comprehensive medical summary reports and answer your health-related queries using the power of Large Language Models by Openai and Langchain.")
    st.markdown("<h2 style='text-align:center;font-family:Georgia'>Features</h1>", unsafe_allow_html=True)
    st.markdown(" - 🤖 AIVABOT MedicalGPT - This Bot is ready to answer your health-related questions")
    st.markdown(" - 📋 Medical Summary Report Generator - Generate detailed medical summary reports for patients")
    st.markdown("-------")
    openai_api_key = st.text_input('Enter OpenAI API Key', type='password')
    st.markdown("-------")
    st.markdown("<h1 style='text-align:center;font-family:Georgia'>📋 Medical Summary Report Generator</h1>", unsafe_allow_html=True)
    patient_name = st.text_input("Patient's Name")
    patient_age = st.number_input("Patient's Age", min_value=0, max_value=150, value=30)
    patient_gender = st.radio("Patient's Gender", ["Male", "Female", "Other"])
    symptoms = st.text_area("Symptoms (comma-separated)", "Fever, Cough, Fatigue")
    medical_history = st.text_area("Medical History")
    physical_examination = st.text_area("Physical Examination Findings")
    provisional_diagnosis = st.text_area("Provisional Diagnosis")
    investigations = st.text_area("Investigations Conducted")
    treatment_plan = st.text_area("Treatment Plan")
    follow_up_instructions = st.text_area("Follow-Up Instructions")
    st.markdown("-------")

    generatebutt = st.button("Generate Medical Summary Report")

def generate_report(patient_name, patient_age, patient_gender, symptoms, medical_history, physical_examination, 
                    provisional_diagnosis, investigations, treatment_plan, follow_up_instructions, report_date):
    doc = docx.Document()

    # Add Title Page followed by section summary
    doc.add_heading("Medical Summary Report", 0)
    doc.add_paragraph(f'Authored By: AIVABOT MedicalGPT LLM')
    doc.add_paragraph(f'Created On: {str(report_date)}')
    doc.add_paragraph(f'Patient: {patient_name}, Age: {patient_age}, Gender: {patient_gender}')
    doc.add_heading(f'Medical Summary for {patient_name}')
    
    # Patient Information
    doc.add_heading('Patient Information')
    doc.add_paragraph(f'Name: {patient_name}')
    doc.add_paragraph(f'Age: {patient_age}')
    doc.add_paragraph(f'Gender: {patient_gender}')

    # Symptoms
    doc.add_heading('Symptoms')
    doc.add_paragraph(symptoms)

    # Medical History
    doc.add_heading('Medical History')
    doc.add_paragraph(medical_history)

    # Physical Examination Findings
    doc.add_heading('Physical Examination Findings')
    doc.add_paragraph(physical_examination)

    # Provisional Diagnosis
    doc.add_heading('Provisional Diagnosis')
    doc.add_paragraph(provisional_diagnosis)

    # Investigations Conducted
    doc.add_heading('Investigations Conducted')
    doc.add_paragraph(investigations)

    # Treatment Plan
    doc.add_heading('Treatment Plan')
    doc.add_paragraph(treatment_plan)

    # Follow-Up Instructions
    doc.add_heading('Follow-Up Instructions')
    doc.add_paragraph(follow_up_instructions)

    doc.save('Medical Summary Report.docx')
    data = open('Medical Summary Report.docx', "rb").read()
    encoded = base64.b64encode(data)
    decoded = base64.b64decode(encoded)

    st.download_button('Download Medical Report Here', decoded, "Medical Summary Report.docx")

def generate_response(input_text):
    llm = OpenAI(temperature=0.3, openai_api_key=openai_api_key)
    output = llm(input_text)
    return output

# App framework
st.markdown("<h1 style='text-align:justified;font-family:Georgia'>🤖 Chatbot</h1>", unsafe_allow_html=True)
# User input field
# Initialize chat history
if "messages" not in st.session_state:
    st.session_state.messages = []
    st.session_state.messages.append({"role": "🤖", "content": "Hello! I'm AIVABOT, here to assist you with medical queries. Please enter your API Key on the sidebar before we get started."})
# Display chat messages from history on app rerun
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        with st.spinner('Starting Bot ...'):
            st.markdown(message["content"])

if generatebutt:
    if openai_api_key.startswith('sk-'):
        date_today = datetime.date.today()
        report_generated_summary = generate_response(f"I need assistance in generating a detailed medical summary report for my patient named {patient_name}. The patient is {patient_age} years old, and the symptoms include {symptoms}. The medical history is {medical_history}, and the provisional diagnosis is {provisional_diagnosis}. Physical examination findings are {physical_examination}. Investigations conducted include {investigations}. The treatment plan is {treatment_plan}, and follow-up instructions are {follow_up_instructions}. Please provide a comprehensive report with all relevant details.")
        generate_report(patient_name, patient_age, patient_gender, symptoms, medical_history, physical_examination, 
                        provisional_diagnosis, investigations, treatment_plan, follow_up_instructions, date_today)
    else:
        st.warning('Please enter your OpenAI API key!', icon='⚠')

if (prompt := st.chat_input("What can I help you with today?")):
    if openai_api_key.startswith('sk-'):
        # Display user message in chat message container
        with st.chat_message("🙂"):
            st.markdown(prompt)
        # Add user message to chat history
        st.session_state.messages.append({"role": "🙂", "content": prompt})

        # Display assistant response in chat message container
        with st.chat_message("🤖"):
            message_placeholder = st.empty()
            full_response = ""
            with st.spinner('Wait for it...'):
                assistant_response = generate_response(prompt)

        for chunk in assistant_response.split():
            full_response += chunk + " "
            time.sleep(0.05)
            # Add a blinking cursor to simulate typing
            message_placeholder.markdown(full_response + "▌")
        message_placeholder.markdown(full_response)

        # Add assistant response to chat history
        st.session_state.messages.append({"role": "🤖", "content": full_response})
    else:
        st.warning('Please enter your OpenAI API key!', icon='⚠')
